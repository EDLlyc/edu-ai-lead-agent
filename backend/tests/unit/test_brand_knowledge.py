from __future__ import annotations

from datetime import date
from io import BytesIO
from pathlib import Path
from uuid import UUID, uuid4
from zipfile import ZIP_DEFLATED, ZipFile

import app.infrastructure.brand.parser as brand_parser_module
import pytest
from app.core.errors import BrandUploadRejectedError
from app.domain.brand_knowledge import (
    BrandAudience,
    BrandClaimScope,
    BrandContentType,
    BrandDocumentKind,
    BrandRetrievalHit,
    BrandSectionKind,
    BrandUploadMetadata,
    ParsedBrandDocument,
    ParsedBrandSection,
    classify_brand_chunk,
    sanitize_brand_filename,
    validated_brand_upload,
)
from app.domain.brand_retrieval import RankedBrandHit, select_diverse_brand_hits
from app.infrastructure.brand.parser import BoundedBrandDocumentParser
from app.schemas.brand_knowledge import BrandContextResponse, BrandRetrievalRequest
from docx import Document as DocxDocument
from pypdf import PdfWriter

FIXTURE = Path(__file__).parents[1] / "fixtures" / "brand" / "parent-tone-v1.md"


def _parser(*, max_characters: int = 20_000) -> BoundedBrandDocumentParser:
    return BoundedBrandDocumentParser(
        max_pages=20,
        max_characters=max_characters,
        max_chunks=50,
        chunk_characters=120,
        overlap_characters=20,
        parser_version="brand-parser-v3-source-structure",
        chunk_version="brand-chunk-v3-parent-child",
        embedding_input_version="brand-embedding-input-v2-section-context",
    )


def test_markdown_validation_parse_and_chunk_are_deterministic() -> None:
    body = FIXTURE.read_bytes()
    upload = validated_brand_upload(
        filename="../品牌资料/parent-tone-v1.md",
        declared_media_type="text/markdown",
        body=body,
    )
    parser = _parser()
    parsed = parser.parse(body=upload.body, media_type=upload.media_type)
    version_id = uuid4()

    first = parser.chunk(version_id=version_id, document=parsed)
    second = parser.chunk(version_id=version_id, document=parsed)

    assert upload.safe_filename == "parent-tone-v1.md"
    assert first == second
    assert len(first) >= 2
    assert all(chunk.char_end <= len(parsed.text) for chunk in first)
    assert all(parsed.text[chunk.char_start : chunk.char_end] == chunk.text for chunk in first)
    assert "品牌材料只能指导表达方式" in parsed.text


def test_sparse_pdf_returns_typed_ocr_handoff() -> None:
    buffer = BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.add_blank_page(width=612, height=792)
    writer.write(buffer)

    parsed = _parser().parse(body=buffer.getvalue(), media_type="application/pdf")

    assert parsed.text == ""
    assert parsed.page_count == 2
    assert parsed.extraction_method == "local"
    assert parsed.requires_ocr is True
    assert parsed.ocr_required is True


def test_text_bearing_pdf_keeps_local_fast_path(monkeypatch: pytest.MonkeyPatch) -> None:
    parser = _parser()
    monkeypatch.setattr(
        parser,
        "_parse_pdf",
        lambda _: ParsedBrandDocument(text="可检索的品牌定位内容。", page_count=1),
    )

    parsed = parser.parse(body=b"ignored", media_type="application/pdf")

    assert parsed.text == "可检索的品牌定位内容。"
    assert parsed.extraction_method == "local"
    assert parsed.requires_ocr is False


def test_chunk_offsets_match_trimmed_text_exactly() -> None:
    parser = _parser()
    parsed = ParsedBrandDocument(
        text="第一段结尾。   \n\n   第二段内容用于验证偏移。   ",
        page_count=None,
    )

    chunks = parser.chunk(version_id=uuid4(), document=parsed)

    assert chunks
    assert all(parsed.text[chunk.char_start : chunk.char_end] == chunk.text for chunk in chunks)
    assert all(chunk.text == chunk.text.strip() for chunk in chunks)


def test_ocr_generic_tiny_blocks_coalesce_under_existing_hard_cap() -> None:
    parser = BoundedBrandDocumentParser(
        max_pages=100,
        max_characters=300_000,
        max_chunks=600,
        chunk_characters=300,
        overlap_characters=40,
        parser_version="brand-parser-v3-source-structure",
        chunk_version="brand-chunk-v3-parent-child",
        embedding_input_version="brand-embedding-input-v2-section-context",
    )
    text = "\n\n".join(f"合成微块 {index:04d}。" for index in range(701))
    parsed = ParsedBrandDocument(
        text=text,
        page_count=50,
        extraction_method="ocr",
    )
    version_id = UUID("aaaaaaaa-bbbb-cccc-dddd-eeeeeeeeeeee")

    first = parser.chunk(
        version_id=version_id,
        document=parsed,
        document_title="合成 OCR 文档",
    )
    second = parser.chunk(
        version_id=version_id,
        document=parsed,
        document_title="合成 OCR 文档",
    )

    assert first == second
    assert len(first.chunks) < 600
    assert len(first.chunks) < 701
    assert all(len(chunk.text) <= 300 for chunk in first.chunks)
    assert all(parsed.text[chunk.char_start : chunk.char_end] == chunk.text for chunk in first)
    assert all(
        left.char_end <= right.char_start
        for left, right in zip(first.chunks, first.chunks[1:], strict=False)
    )
    covered = [False] * len(parsed.text)
    for chunk in first.chunks:
        for offset in range(chunk.char_start, chunk.char_end):
            covered[offset] = True
    assert all(character.isspace() or covered[offset] for offset, character in enumerate(text))


def test_ocr_generic_continuous_parent_fallback_preserves_all_content() -> None:
    parser = BoundedBrandDocumentParser(
        max_pages=100,
        max_characters=20_000,
        max_chunks=2,
        chunk_characters=300,
        overlap_characters=40,
        parser_version="brand-parser-v3-source-structure",
        chunk_version="brand-chunk-v3-parent-child",
        embedding_input_version="brand-embedding-input-v2-section-context",
    )
    # Complete 150-character blocks cannot be paired under the 300-character child limit,
    # while a continuous parent-local split can preserve them in two overlapping children.
    text = "\n\n".join(character * 150 for character in ("甲", "乙", "丙"))
    parsed = ParsedBrandDocument(
        text=text,
        page_count=1,
        extraction_method="ocr",
    )
    version_id = UUID("bbbbbbbb-cccc-dddd-eeee-ffffffffffff")

    first = parser.chunk(version_id=version_id, document=parsed)
    second = parser.chunk(version_id=version_id, document=parsed)

    assert first == second
    assert len(first.chunks) == 2
    assert all(len(chunk.text) <= 300 for chunk in first)
    assert all(parsed.text[chunk.char_start : chunk.char_end] == chunk.text for chunk in first)
    covered = [False] * len(text)
    for chunk in first:
        for offset in range(chunk.char_start, chunk.char_end):
            covered[offset] = True
    assert all(character.isspace() or covered[offset] for offset, character in enumerate(text))


@pytest.mark.parametrize(
    ("extraction_method", "section_kind"),
    [
        ("local", BrandSectionKind.GENERIC),
        ("ocr", BrandSectionKind.PAGE),
    ],
)
def test_ocr_budget_fallback_is_restricted_to_generic_ocr_parents(
    extraction_method: str,
    section_kind: BrandSectionKind,
) -> None:
    parser = BoundedBrandDocumentParser(
        max_pages=100,
        max_characters=20_000,
        max_chunks=2,
        chunk_characters=300,
        overlap_characters=40,
        parser_version="brand-parser-v3-source-structure",
        chunk_version="brand-chunk-v3-parent-child",
        embedding_input_version="brand-embedding-input-v2-section-context",
    )
    text = "\n\n".join(character * 150 for character in ("甲", "乙", "丙"))
    parsed = ParsedBrandDocument(
        text=text,
        page_count=1,
        extraction_method=extraction_method,
        sections=(
            ParsedBrandSection(
                ordinal=0,
                kind=section_kind,
                title="合成父级",
                text=text,
                char_start=0,
                char_end=len(text),
                source_page=1 if section_kind == BrandSectionKind.PAGE else None,
            ),
        ),
    )

    with pytest.raises(BrandUploadRejectedError) as error:
        parser.chunk(version_id=uuid4(), document=parsed)

    assert error.value.code == "brand_chunk_limit"


def test_ocr_generic_coalescing_never_crosses_parent_sections() -> None:
    parser = BoundedBrandDocumentParser(
        max_pages=100,
        max_characters=300_000,
        max_chunks=600,
        chunk_characters=300,
        overlap_characters=40,
        parser_version="brand-parser-v3-source-structure",
        chunk_version="brand-chunk-v3-parent-child",
        embedding_input_version="brand-embedding-input-v2-section-context",
    )
    first_parent = "\n\n".join(f"父级甲微块 {index:04d}。" for index in range(620))
    second_parent = "\n\n".join(f"父级乙微块 {index:04d}。" for index in range(620))
    text = f"{first_parent}\n\n{second_parent}"
    second_start = len(first_parent) + 2
    parsed = ParsedBrandDocument(
        text=text,
        page_count=50,
        extraction_method="ocr",
        sections=(
            ParsedBrandSection(
                ordinal=0,
                kind=BrandSectionKind.GENERIC,
                title="合成父级甲",
                text=first_parent,
                char_start=0,
                char_end=len(first_parent),
            ),
            ParsedBrandSection(
                ordinal=1,
                kind=BrandSectionKind.GENERIC,
                title="合成父级乙",
                text=second_parent,
                char_start=second_start,
                char_end=len(text),
            ),
        ),
    )

    result = parser.chunk(version_id=uuid4(), document=parsed)

    assert len(result.chunks) <= 600
    section_by_id = {section.id: section for section in result.sections}
    assert {chunk.section_id for chunk in result.chunks} == set(section_by_id)
    assert all(
        section_by_id[chunk.section_id].char_start <= chunk.char_start
        and chunk.char_end <= section_by_id[chunk.section_id].char_end
        and parsed.text[chunk.char_start : chunk.char_end] == chunk.text
        for chunk in result.chunks
    )
    assert all(("父级甲" in chunk.text) != ("父级乙" in chunk.text) for chunk in result.chunks)


def test_ocr_generic_still_rejects_content_that_cannot_fit_hard_cap() -> None:
    parser = BoundedBrandDocumentParser(
        max_pages=100,
        max_characters=20_000,
        max_chunks=2,
        chunk_characters=300,
        overlap_characters=40,
        parser_version="brand-parser-v3-source-structure",
        chunk_version="brand-chunk-v3-parent-child",
        embedding_input_version="brand-embedding-input-v2-section-context",
    )
    parsed = ParsedBrandDocument(
        text="合" * 1_000,
        page_count=1,
        extraction_method="ocr",
    )

    with pytest.raises(BrandUploadRejectedError) as error:
        parser.chunk(version_id=uuid4(), document=parsed)

    assert error.value.code == "brand_chunk_limit"


def test_chunk_prefers_markdown_paragraph_and_sentence_boundaries() -> None:
    parser = BoundedBrandDocumentParser(
        max_pages=20,
        max_characters=20_000,
        max_chunks=50,
        chunk_characters=48,
        overlap_characters=8,
        parser_version="brand-parser-v2-glm-ocr",
        chunk_version="brand-chunk-v2-structure-aware",
        embedding_input_version="brand-embedding-input-v1",
    )
    parsed = ParsedBrandDocument(
        text=(
            "## 品牌定位\n\n"
            "第一段用于验证结构边界。第二句继续说明表达原则。\n\n"
            "第二段不应被第一片段提前吞入。第三句保持完整。"
        ),
        page_count=None,
    )

    chunks = parser.chunk(version_id=uuid4(), document=parsed)

    assert chunks
    assert len(chunks[0].text) <= 48
    assert chunks[0].text.endswith("。")
    assert "第二段" not in chunks[0].text
    assert all(parsed.text[chunk.char_start : chunk.char_end] == chunk.text for chunk in chunks)


def test_chunk_uses_sentence_boundary_when_no_block_boundary_fits() -> None:
    parser = BoundedBrandDocumentParser(
        max_pages=20,
        max_characters=20_000,
        max_chunks=50,
        chunk_characters=25,
        overlap_characters=8,
        parser_version="brand-parser-v2-glm-ocr",
        chunk_version="brand-chunk-v2-structure-aware",
        embedding_input_version="brand-embedding-input-v1",
    )
    parsed = ParsedBrandDocument(
        text="第一句保持完整。第二句也保持完整。第三句继续说明。第四句用于超过上限。",
        page_count=None,
    )

    chunks = parser.chunk(version_id=uuid4(), document=parsed)

    assert len(chunks) >= 2
    assert chunks[0].text.endswith("。")
    assert all(len(chunk.text) <= 25 for chunk in chunks)
    assert all(parsed.text[chunk.char_start : chunk.char_end] == chunk.text for chunk in chunks)


def _ranked_brand_hit(
    document_id,
    version_id,
    ordinal: int,
    score: float,
    *,
    section_id=None,
) -> RankedBrandHit:
    return RankedBrandHit(
        hit=BrandRetrievalHit(
            chunk_id=uuid4(),
            document_id=document_id,
            version_id=version_id,
            document_title="品牌资料",
            document_kind=BrandDocumentKind.TONE,
            audience=BrandAudience.PARENTS,
            text=f"品牌原则 {document_id} {ordinal}",
            tone_tags=(),
            safety_tags=(),
            visual_tags=(),
            full_text_score=score,
            vector_score=score,
            fused_score=score,
            section_id=section_id,
        ),
        ordinal=ordinal,
    )


def test_retrieval_diversity_prefers_distinct_parents_then_preserves_rrf_order() -> None:
    document_a, document_b, document_c = uuid4(), uuid4(), uuid4()
    version_a, version_b, version_c = uuid4(), uuid4(), uuid4()
    section_a, section_b, section_c = uuid4(), uuid4(), uuid4()
    candidates = [
        _ranked_brand_hit(document_a, version_a, 0, 0.90, section_id=section_a),
        _ranked_brand_hit(document_a, version_a, 1, 0.89, section_id=section_a),
        _ranked_brand_hit(document_a, version_a, 2, 0.88, section_id=uuid4()),
        _ranked_brand_hit(document_b, version_b, 0, 0.87, section_id=section_b),
        _ranked_brand_hit(document_c, version_c, 0, 0.86, section_id=section_c),
    ]

    selected = select_diverse_brand_hits(candidates, limit=4)

    assert [hit.fused_score for hit in selected] == [0.90, 0.88, 0.87, 0.86]
    assert [hit.document_id for hit in selected] == [document_a, document_a, document_b, document_c]
    assert len({hit.document_id for hit in selected}) == 3
    ordinal_by_chunk_id = {candidate.hit.chunk_id: candidate.ordinal for candidate in candidates}
    assert [ordinal_by_chunk_id[hit.chunk_id] for hit in selected] == [0, 2, 0, 0]


def test_pdf_page_card_blocks_become_parent_local_children() -> None:
    parser = _parser()
    text = "页面主题\n\n卡片甲介绍产品能力。\n\n卡片乙介绍安全能力。"
    parsed = ParsedBrandDocument(
        text=text,
        page_count=1,
        sections=(
            ParsedBrandSection(
                ordinal=0,
                kind=BrandSectionKind.PAGE,
                title="页面主题",
                text=text,
                char_start=0,
                char_end=len(text),
                source_page=1,
            ),
        ),
    )

    result = parser.chunk(
        version_id=uuid4(),
        document=parsed,
        document_title="合成平台介绍",
        document_kind=BrandDocumentKind.OTHER,
    )

    assert len(result.sections) == 1
    assert [chunk.section_ordinal for chunk in result.chunks] == [0, 1]
    assert "卡片甲" in result.chunks[0].text
    assert result.chunks[1].text == "卡片乙介绍安全能力。"
    assert all(chunk.section_id == result.sections[0].id for chunk in result.chunks)
    assert all(parsed.text[chunk.char_start : chunk.char_end] == chunk.text for chunk in result)


def test_pdf_parse_preserves_nonempty_source_pages_and_stable_offsets(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    class _Page:
        def __init__(self, text: str) -> None:
            self._text = text

        def extract_text(self) -> str:
            return self._text

    class _Reader:
        is_encrypted = False

        def __init__(self, *_: object, **__: object) -> None:
            self.pages = [
                _Page("第一页主题\n\n卡片甲说明。" * 5),
                _Page(""),
                _Page("第三页主题\n\n卡片乙说明。" * 5),
            ]

    monkeypatch.setattr(brand_parser_module, "PdfReader", _Reader)
    parser = _parser()
    parsed = parser.parse(body=b"synthetic", media_type="application/pdf")
    version_id = uuid4()

    first = parser.chunk(
        version_id=version_id,
        document=parsed,
        document_title="合成演示稿",
        document_kind=BrandDocumentKind.OTHER,
    )
    second = parser.chunk(
        version_id=version_id,
        document=parsed,
        document_title="合成演示稿",
        document_kind=BrandDocumentKind.OTHER,
    )

    assert parsed.page_count == 3
    assert [section.source_page for section in parsed.sections] == [1, 3]
    assert first == second
    section_by_id = {section.id: section for section in first.sections}
    assert all(
        section_by_id[chunk.section_id].char_start <= chunk.char_start
        and chunk.char_end <= section_by_id[chunk.section_id].char_end
        for chunk in first
    )


def test_frozen_v2_and_v3_derivation_bundles_have_distinct_literal_fingerprints(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    class _Page:
        def __init__(self, text: str) -> None:
            self._text = text

        def extract_text(self) -> str:
            return self._text

    class _Reader:
        is_encrypted = False

        def __init__(self, *_: object, **__: object) -> None:
            self.pages = [
                _Page("第一页主题\n\n合成卡片甲说明。"),
                _Page("第二页主题\n\n合成卡片乙说明。"),
            ]

    monkeypatch.setattr(brand_parser_module, "PdfReader", _Reader)
    common = {
        "max_pages": 20,
        "max_characters": 20_000,
        "max_chunks": 50,
        "chunk_characters": 300,
        "overlap_characters": 20,
    }
    legacy = BoundedBrandDocumentParser(
        **common,
        parser_version="brand-parser-v2-glm-ocr",
        chunk_version="brand-chunk-v2-structure-aware",
        embedding_input_version="brand-embedding-input-v1",
    )
    structured = BoundedBrandDocumentParser(
        **common,
        parser_version="brand-parser-v3-source-structure",
        chunk_version="brand-chunk-v3-parent-child",
        embedding_input_version="brand-embedding-input-v2-section-context",
    )
    version_id = UUID("11111111-2222-3333-4444-555555555555")

    legacy_document = legacy.parse(body=b"synthetic", media_type="application/pdf")
    legacy_result = legacy.chunk(version_id=version_id, document=legacy_document)
    structured_document = structured.parse(body=b"synthetic", media_type="application/pdf")
    structured_result = structured.chunk(
        version_id=version_id,
        document=structured_document,
        document_title="合成演示稿",
        document_kind=BrandDocumentKind.OTHER,
    )

    assert legacy_document.sections == ()
    assert legacy_result.sections == ()
    assert len(legacy_result.chunks) == 1
    assert legacy_result.chunks[0].section_id is None
    assert legacy_result.chunks[0].embedding_text == legacy_result.chunks[0].text
    assert legacy_result.chunks[0].embedding_input_hash == legacy_result.chunks[0].text_hash
    assert legacy_result.chunks[0].chunk_key == (
        "28c1059ba28e935a34bb3857ee6aa0711a9ac7990b5e045a93785153c43cebec"
    )
    assert legacy_result.chunks[0].embedding_input_hash == (
        "07c1c1b90e27d652606217fe77741c59d589cf28592601ba96fe395a258774cf"
    )
    assert [section.source_page for section in structured_result.sections] == [1, 2]
    assert len(structured_result.chunks) == 2
    assert structured_result.chunks[0].chunk_key == (
        "e8eaad6417cf71c8b897ec8505434d5d7fd26a94fe215c3f3185ddcec110cf5f"
    )
    assert structured_result.chunks[0].embedding_input_hash == (
        "f2bdb2c9ebb10e85e0ba67d4d4e19c8a8f1793453785c3e22b223dd7b50cfda6"
    )


def test_parser_rejects_mixed_derivation_versions() -> None:
    with pytest.raises(ValueError, match="unsupported brand parser/chunk/input version bundle"):
        BoundedBrandDocumentParser(
            max_pages=20,
            max_characters=20_000,
            max_chunks=50,
            chunk_characters=300,
            overlap_characters=20,
            parser_version="brand-parser-v2-glm-ocr",
            chunk_version="brand-chunk-v3-parent-child",
            embedding_input_version="brand-embedding-input-v1",
        )


def test_docx_interview_preserves_table_order_and_question_context() -> None:
    document = DocxDocument()
    document.add_heading("访谈摘录", level=1)
    document.add_paragraph("Q1. 品牌为什么重视探索\uff1f")
    document.add_paragraph("第一段回答说明探索价值。")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "观察"
    table.cell(0, 1).text = "实践"
    document.add_paragraph("Q2\uff1a平台怎样支持学习\uff1f")
    document.add_paragraph("课程通过真实问题支持学习。" * 12)
    buffer = BytesIO()
    document.save(buffer)

    parser = _parser()
    parsed = parser.parse(
        body=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    result = parser.chunk(
        version_id=uuid4(),
        document=parsed,
        document_title="合成访谈",
        document_kind=BrandDocumentKind.POSITIONING,
    )

    questions = [
        section for section in parsed.sections if section.kind == BrandSectionKind.INTERVIEW_QA
    ]
    assert [section.question_number for section in questions] == [1, 2]
    assert parsed.text.index("观察 实践") < parsed.text.index("Q2")
    second_parent = result.sections[questions[1].ordinal]
    second_children = [chunk for chunk in result if chunk.section_id == second_parent.id]
    assert len(second_children) >= 2
    assert all("问题\uff1a平台怎样支持学习?" in chunk.embedding_text for chunk in second_children)
    assert all(parsed.text[chunk.char_start : chunk.char_end] == chunk.text for chunk in result)

    legacy_parser = BoundedBrandDocumentParser(
        max_pages=20,
        max_characters=20_000,
        max_chunks=50,
        chunk_characters=900,
        overlap_characters=120,
        parser_version="brand-parser-v2-glm-ocr",
        chunk_version="brand-chunk-v2-structure-aware",
        embedding_input_version="brand-embedding-input-v1",
    )
    legacy_parsed = legacy_parser.parse(
        body=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert legacy_parsed.sections == ()
    assert legacy_parsed.text.index("Q2") < legacy_parsed.text.index("观察 实践")


@pytest.mark.parametrize(
    ("title", "text", "expected_type"),
    [
        ("品牌定位", "面向青少年提供学习支持。", BrandContentType.POSITIONING),
        ("课程产品", "适用年龄为小学阶段。", BrandContentType.PRODUCT_PROFILE),
        ("家长需求", "帮助理解用户痛点。", BrandContentType.AUDIENCE_INSIGHT),
        ("安全护栏", "内容过滤保护使用安全。", BrandContentType.SAFETY_CAPABILITY),
        ("数字 IP 价值", "角色价值保持一致。", BrandContentType.DIGITAL_IP_VALUES),
        ("表达方式", "品牌语气温暖准确。", BrandContentType.TONE_EXAMPLE),
        ("视觉规范", "插画色彩保持统一。", BrandContentType.VISUAL_GUIDANCE),
        ("其他内容", "没有可确定的类别。", BrandContentType.OTHER),
    ],
)
def test_chunk_content_classification_is_deterministic(
    title: str, text: str, expected_type: BrandContentType
) -> None:
    content_type, claim_scope, verification = classify_brand_chunk(
        document_kind=BrandDocumentKind.OTHER,
        section_title=title,
        question_text=None,
        text=text,
    )

    assert content_type == expected_type
    assert claim_scope == BrandClaimScope.BRAND_STATEMENT
    assert verification is False


def test_external_claim_overrides_normative_scope_without_erasing_content_type() -> None:
    content_type, claim_scope, verification = classify_brand_chunk(
        document_kind=BrandDocumentKind.SAFETY_RULE,
        section_title="安全能力",
        question_text=None,
        text="必须使用通过认证的安全过滤体系。",
    )

    assert content_type == BrandContentType.SAFETY_CAPABILITY
    assert claim_scope == BrandClaimScope.EXTERNAL_CLAIM
    assert verification is True


def test_untyped_external_claim_receives_external_content_type() -> None:
    content_type, claim_scope, verification = classify_brand_chunk(
        document_kind=BrandDocumentKind.OTHER,
        section_title="第三方信息",
        question_text=None,
        text="该项目获得认证。",
    )

    assert content_type == BrandContentType.EXTERNAL_CLAIM
    assert claim_scope == BrandClaimScope.EXTERNAL_CLAIM
    assert verification is True


@pytest.mark.parametrize(
    "text",
    (
        "项目已完成融资。",
        "品牌获得行业奖项。",
        "项目与第三方机构合作。",
        "课程覆盖一百所学校。",
    ),
)
def test_external_claim_terms_without_arabic_numbers_still_require_verification(
    text: str,
) -> None:
    _, claim_scope, verification = classify_brand_chunk(
        document_kind=BrandDocumentKind.OTHER,
        section_title="项目信息",
        question_text=None,
        text=text,
    )

    assert claim_scope == BrandClaimScope.EXTERNAL_CLAIM
    assert verification is True


def test_retrieval_diversity_falls_back_when_only_one_document_is_available() -> None:
    document_id, version_id = uuid4(), uuid4()
    candidates = [
        _ranked_brand_hit(document_id, version_id, 0, 0.90),
        _ranked_brand_hit(document_id, version_id, 1, 0.89),
        _ranked_brand_hit(document_id, version_id, 2, 0.88),
    ]

    selected = select_diverse_brand_hits(candidates, limit=3)

    assert [hit.fused_score for hit in selected] == [0.90, 0.89, 0.88]
    assert len(selected) == 3


def test_retrieval_diversity_exhausts_unseen_parent_before_second_child() -> None:
    document_id, version_id = uuid4(), uuid4()
    section_a, section_b, section_c = uuid4(), uuid4(), uuid4()
    candidates = [
        _ranked_brand_hit(document_id, version_id, 0, 0.90, section_id=section_a),
        _ranked_brand_hit(document_id, version_id, 1, 0.89, section_id=section_a),
        _ranked_brand_hit(document_id, version_id, 2, 0.88, section_id=section_b),
        _ranked_brand_hit(document_id, version_id, 3, 0.87, section_id=section_c),
    ]

    selected = select_diverse_brand_hits(candidates, limit=3)
    legacy_selected = select_diverse_brand_hits(
        candidates,
        limit=3,
        retrieval_version="brand-hybrid-rrf-v2-diverse",
    )

    assert [hit.section_id for hit in selected] == [section_a, section_b, section_c]
    assert [hit.section_id for hit in legacy_selected] == [section_a, section_a, section_b]


def test_brand_metadata_fingerprint_tracks_semantic_version_metadata() -> None:
    base = BrandUploadMetadata(
        brand_slug="sai-xiansheng",
        title="赛先生家长沟通规范",
        document_kind=BrandDocumentKind.TONE,
        audience=BrandAudience.PARENTS,
        language="zh-CN",
        valid_from=date(2026, 7, 1),
        valid_until=None,
        tone_tags=("温暖", "准确"),
        safety_tags=("不制造焦虑",),
        visual_tags=(),
    )
    reordered = BrandUploadMetadata(
        brand_slug="sai-xiansheng",
        title="赛先生家长沟通规范",
        document_kind=BrandDocumentKind.TONE,
        audience=BrandAudience.PARENTS,
        language="zh-CN",
        valid_from=date(2026, 7, 1),
        valid_until=None,
        tone_tags=("准确", "温暖"),
        safety_tags=("不制造焦虑",),
        visual_tags=(),
    )
    changed = BrandUploadMetadata(
        brand_slug="sai-xiansheng",
        title="赛先生家长沟通规范",
        document_kind=BrandDocumentKind.TONE,
        audience=BrandAudience.PARENTS,
        language="zh-CN",
        valid_from=date(2026, 7, 15),
        valid_until=None,
        tone_tags=("准确", "温暖"),
        safety_tags=("不制造焦虑",),
        visual_tags=(),
    )

    assert base.metadata_fingerprint == reordered.metadata_fingerprint
    assert base.metadata_fingerprint != changed.metadata_fingerprint


def test_upload_rejects_signature_extension_mismatch_and_non_utf8_text() -> None:
    with pytest.raises(ValueError, match="signature"):
        validated_brand_upload(
            filename="rules.pdf",
            declared_media_type="application/pdf",
            body=b"not-a-pdf",
        )
    upload = validated_brand_upload(
        filename="rules.txt",
        declared_media_type="text/plain",
        body=b"\xff\xfeunsafe",
    )
    with pytest.raises(BrandUploadRejectedError, match="UTF-8"):
        _parser().parse(body=upload.body, media_type=upload.media_type)


def test_docx_rejects_external_relationships_and_embedded_objects() -> None:
    buffer = BytesIO()
    with ZipFile(buffer, "w", ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "<Types />")
        archive.writestr(
            "word/_rels/document.xml.rels",
            '<Relationships><Relationship TargetMode="External" Target="https://example.com"/></Relationships>',
        )
        archive.writestr("word/document.xml", "<document />")
    with pytest.raises(BrandUploadRejectedError, match="external relationships"):
        _parser().parse(
            body=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )


def test_parser_enforces_character_limit_without_returning_corpus_text() -> None:
    with pytest.raises(BrandUploadRejectedError) as captured:
        _parser(max_characters=1_000).parse(
            body=("家长沟通原则。" * 500).encode(),
            media_type="text/plain",
        )
    assert captured.value.code == "brand_text_too_large"
    assert "家长沟通原则" not in captured.value.message


def test_filename_is_bounded_and_drops_parent_paths() -> None:
    assert sanitize_brand_filename("../../赛先生 语气规范.md") == "赛先生-语气规范.md"
    with pytest.raises(ValueError):
        sanitize_brand_filename("../..")


def test_retrieval_schema_describes_internal_copy_generation_contract() -> None:
    request_schema = BrandRetrievalRequest.model_json_schema()
    response_schema = BrandContextResponse.model_json_schema()

    assert "copy-generation" in request_schema["description"]
    audience_description = request_schema["properties"]["audience"]["description"]
    assert "target audience" in audience_description.lower()
    assert "not the identity of a search user" in audience_description
    assert "copy generation" in response_schema["description"]
    assert response_schema["properties"]["evidence_eligible"]["const"] is False
