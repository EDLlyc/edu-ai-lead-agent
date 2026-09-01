from __future__ import annotations

from io import BytesIO
from uuid import UUID

import app.infrastructure.brand.parser as brand_parser_module
import pytest
from app.core.errors import BrandUploadRejectedError
from app.domain.brand_knowledge import (
    BrandDocumentKind,
    BrandLayoutSemanticRole,
    BrandOcrBlockKind,
    BrandOcrLayoutBlock,
    BrandOcrLayoutPage,
)
from app.infrastructure.brand.parser import (
    BoundedBrandDocumentParser,
    PdfPageQualityProfile,
    requires_layout_ocr,
)
from docx import Document as DocxDocument


def _parser(*, layout: bool = True) -> BoundedBrandDocumentParser:
    versions = (
        (
            "brand-parser-v4-layout-aware",
            "brand-chunk-v4-layout-blocks",
            "brand-embedding-input-v2-section-context",
        )
        if layout
        else (
            "brand-parser-v3-source-structure",
            "brand-chunk-v3-parent-child",
            "brand-embedding-input-v2-section-context",
        )
    )
    return BoundedBrandDocumentParser(
        max_pages=100,
        max_characters=300_000,
        max_chunks=600,
        chunk_characters=300,
        overlap_characters=40,
        parser_version=versions[0],
        chunk_version=versions[1],
        embedding_input_version=versions[2],
        sparse_text_threshold=40,
    )


@pytest.mark.parametrize(
    ("profile", "expected"),
    (
        (PdfPageQualityProfile(10, 400, 1, 1, 8), False),
        (PdfPageQualityProfile(10, 400, 1, 2, 8), True),
        (PdfPageQualityProfile(10, 400, 1, 2, 7), False),
        (PdfPageQualityProfile(10, 399, 0, 0, 0), True),
    ),
)
def test_layout_route_has_frozen_sparse_and_slide_ratio_edges(
    profile: PdfPageQualityProfile,
    expected: bool,
) -> None:
    assert requires_layout_ocr(profile, sparse_text_threshold=40) is expected


def test_v4_routes_sparse_slide_deck_without_changing_v3(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    class _MediaBox:
        width = 960
        height = 540

    class _Page:
        mediabox = _MediaBox()

        def __init__(self, text: str) -> None:
            self._text = text

        def extract_text(self) -> str:
            return self._text

    class _Reader:
        is_encrypted = False

        def __init__(self, *_: object, **__: object) -> None:
            self.pages = (
                [_Page("")] * 5 + [_Page("稀疏页" * 7)] * 31 + [_Page("密集页面内容" * 22)] * 12
            )

    monkeypatch.setattr(brand_parser_module, "PdfReader", _Reader)

    v4 = _parser().parse(body=b"synthetic", media_type="application/pdf")
    v3 = _parser(layout=False).parse(body=b"synthetic", media_type="application/pdf")

    assert v4.page_count == v3.page_count == 48
    assert v4.requires_ocr is True
    assert v3.requires_ocr is False
    assert [section.source_page for section in v4.sections] == list(range(6, 49))
    assert v4.text == v3.text


def test_v4_layout_pages_preserve_exact_blocks_and_conservative_cards() -> None:
    pages = (
        BrandOcrLayoutPage(page_number=1, blocks=()),
        BrandOcrLayoutPage(
            page_number=2,
            blocks=(
                BrandOcrLayoutBlock(
                    ordinal=0,
                    kind=BrandOcrBlockKind.TEXT,
                    text="安全护栏",
                    semantic_role=BrandLayoutSemanticRole.PARAGRAPH_TITLE,
                    normalized_bbox=(0.08, 0.10, 0.42, 0.16),
                ),
                BrandOcrLayoutBlock(
                    ordinal=1,
                    kind=BrandOcrBlockKind.TEXT,
                    text="过滤与预审共同保护使用安全。",
                    semantic_role=BrandLayoutSemanticRole.CONTENT,
                    normalized_bbox=(0.08, 0.18, 0.42, 0.34),
                ),
                BrandOcrLayoutBlock(
                    ordinal=2,
                    kind=BrandOcrBlockKind.TABLE,
                    text="|能力|说明|\n|---|---|\n|过滤|风险控制|",
                    semantic_role=BrandLayoutSemanticRole.TABLE,
                    normalized_bbox=(0.52, 0.10, 0.92, 0.38),
                ),
            ),
            width=960,
            height=540,
        ),
        BrandOcrLayoutPage(
            page_number=3,
            blocks=(
                BrandOcrLayoutBlock(
                    ordinal=0,
                    kind=BrandOcrBlockKind.FORMULA,
                    text="E = mc^2",
                    semantic_role=BrandLayoutSemanticRole.DISPLAY_FORMULA,
                ),
            ),
        ),
    )
    parser = _parser()
    parsed = parser.parse_ocr(markdown="compatibility only", layout_pages=pages, page_count=3)
    version_id = UUID("11111111-2222-4333-8444-555555555555")

    first = parser.chunk(
        version_id=version_id,
        document=parsed,
        document_title="合成演示稿",
        document_kind=BrandDocumentKind.OTHER,
    )
    second = parser.chunk(
        version_id=version_id,
        document=parsed,
        document_title="合成演示稿",
        document_kind=BrandDocumentKind.OTHER,
    )

    assert first == second
    assert [section.source_page for section in parsed.sections] == [2, 3]
    assert parsed.sections[0].title == "安全护栏"
    assert [block.semantic_role for block in parsed.sections[0].layout_blocks] == [
        BrandLayoutSemanticRole.PARAGRAPH_TITLE,
        BrandLayoutSemanticRole.CONTENT,
        BrandLayoutSemanticRole.TABLE,
    ]
    assert [chunk.text for chunk in first.chunks] == [
        "安全护栏\n\n过滤与预审共同保护使用安全。",
        "|能力|说明|\n|---|---|\n|过滤|风险控制|",
        "E = mc^2",
    ]
    assert all(parsed.text[chunk.char_start : chunk.char_end] == chunk.text for chunk in first)
    assert all(
        parsed.text[block.char_start : block.char_end] == block.text
        for section in parsed.sections
        for block in section.layout_blocks
    )


def test_v4_prefers_explicit_semantic_title_over_topmost_text() -> None:
    pages = (
        BrandOcrLayoutPage(
            page_number=1,
            blocks=(
                BrandOcrLayoutBlock(
                    ordinal=0,
                    kind=BrandOcrBlockKind.TEXT,
                    text="页眉说明",
                    semantic_role=BrandLayoutSemanticRole.CONTENT,
                    normalized_bbox=(0.08, 0.02, 0.92, 0.07),
                ),
                BrandOcrLayoutBlock(
                    ordinal=1,
                    kind=BrandOcrBlockKind.TEXT,
                    text="明确页面标题",
                    semantic_role=BrandLayoutSemanticRole.DOC_TITLE,
                    normalized_bbox=(0.08, 0.12, 0.60, 0.20),
                ),
                BrandOcrLayoutBlock(
                    ordinal=2,
                    kind=BrandOcrBlockKind.TEXT,
                    text="页面正文内容。",
                    semantic_role=BrandLayoutSemanticRole.CONTENT,
                    normalized_bbox=(0.08, 0.23, 0.60, 0.36),
                ),
            ),
        ),
    )

    parsed = _parser().parse_ocr(
        markdown="compatibility only",
        layout_pages=pages,
        page_count=1,
    )

    assert parsed.sections[0].title == "明确页面标题"
    assert all(
        parsed.text[block.char_start : block.char_end] == block.text
        for block in parsed.sections[0].layout_blocks
    )


def test_v4_never_promotes_captions_or_footnotes_to_content_titles() -> None:
    pages = (
        BrandOcrLayoutPage(
            page_number=1,
            blocks=(
                BrandOcrLayoutBlock(
                    ordinal=0,
                    kind=BrandOcrBlockKind.TEXT,
                    text="图 1 | 安全能力",
                    semantic_role=BrandLayoutSemanticRole.FIGURE_TITLE,
                    normalized_bbox=(0.08, 0.10, 0.70, 0.16),
                ),
                BrandOcrLayoutBlock(
                    ordinal=1,
                    kind=BrandOcrBlockKind.TEXT,
                    text="图中展示过滤流程。",
                    semantic_role=BrandLayoutSemanticRole.CONTENT,
                    normalized_bbox=(0.08, 0.18, 0.70, 0.32),
                ),
            ),
        ),
        BrandOcrLayoutPage(
            page_number=2,
            blocks=(
                BrandOcrLayoutBlock(
                    ordinal=0,
                    kind=BrandOcrBlockKind.TEXT,
                    text="资料来源说明",
                    semantic_role=BrandLayoutSemanticRole.VISION_FOOTNOTE,
                    normalized_bbox=(0.08, 0.78, 0.70, 0.84),
                ),
                BrandOcrLayoutBlock(
                    ordinal=1,
                    kind=BrandOcrBlockKind.TEXT,
                    text="脚注后的独立正文。",
                    semantic_role=BrandLayoutSemanticRole.CONTENT,
                    normalized_bbox=(0.08, 0.86, 0.70, 0.96),
                ),
            ),
        ),
        BrandOcrLayoutPage(
            page_number=3,
            blocks=(
                BrandOcrLayoutBlock(
                    ordinal=0,
                    kind=BrandOcrBlockKind.TEXT,
                    text="通用文本标题",
                    semantic_role=BrandLayoutSemanticRole.TEXT,
                    normalized_bbox=(0.08, 0.10, 0.70, 0.16),
                ),
                BrandOcrLayoutBlock(
                    ordinal=1,
                    kind=BrandOcrBlockKind.TEXT,
                    text="通用文本仍按几何规则合并。",
                    semantic_role=BrandLayoutSemanticRole.CONTENT,
                    normalized_bbox=(0.08, 0.18, 0.70, 0.32),
                ),
            ),
        ),
        BrandOcrLayoutPage(
            page_number=4,
            blocks=(
                BrandOcrLayoutBlock(
                    ordinal=0,
                    kind=BrandOcrBlockKind.TEXT,
                    text="省略原生角色",
                    normalized_bbox=(0.08, 0.10, 0.70, 0.16),
                ),
                BrandOcrLayoutBlock(
                    ordinal=1,
                    kind=BrandOcrBlockKind.TEXT,
                    text="兼容既有几何合并行为。",
                    normalized_bbox=(0.08, 0.18, 0.70, 0.32),
                ),
            ),
        ),
    )
    parser = _parser()
    parsed = parser.parse_ocr(
        markdown="compatibility only",
        layout_pages=pages,
        page_count=4,
    )
    result = parser.chunk(version_id=UUID(int=8), document=parsed)

    assert [section.title for section in parsed.sections] == [
        "第 1 页",
        "第 2 页",
        "通用文本标题",
        "省略原生角色",
    ]
    assert [chunk.text for chunk in result.chunks] == [
        "图 1 | 安全能力",
        "图中展示过滤流程。",
        "资料来源说明",
        "脚注后的独立正文。",
        "通用文本标题\n\n通用文本仍按几何规则合并。",
        "省略原生角色\n\n兼容既有几何合并行为。",
    ]


def test_v4_semantic_role_groups_are_exhaustive_and_disjoint() -> None:
    title_roles = brand_parser_module._LAYOUT_PAGE_TITLE_SEMANTIC_ROLES
    generic_roles = brand_parser_module._LAYOUT_GENERIC_TEXT_SEMANTIC_ROLES
    body_only_roles = brand_parser_module._LAYOUT_CARD_BODY_SEMANTIC_ROLES - generic_roles
    non_title_non_card_roles = brand_parser_module._LAYOUT_NON_TITLE_NON_CARD_SEMANTIC_ROLES
    groups = (title_roles, generic_roles, body_only_roles, non_title_non_card_roles)

    assert frozenset().union(*groups) == frozenset(BrandLayoutSemanticRole)
    assert sum(len(group) for group in groups) == len(BrandLayoutSemanticRole)
    assert {
        BrandLayoutSemanticRole.ASIDE_TEXT,
        BrandLayoutSemanticRole.FOOTER,
        BrandLayoutSemanticRole.FOOTER_IMAGE,
        BrandLayoutSemanticRole.FOOTNOTE,
        BrandLayoutSemanticRole.HEADER,
        BrandLayoutSemanticRole.HEADER_IMAGE,
        BrandLayoutSemanticRole.NUMBER,
        BrandLayoutSemanticRole.REFERENCE,
    } <= non_title_non_card_roles


@pytest.mark.parametrize(
    "semantic_role",
    (
        BrandLayoutSemanticRole.ASIDE_TEXT,
        BrandLayoutSemanticRole.FOOTER,
        BrandLayoutSemanticRole.FOOTER_IMAGE,
        BrandLayoutSemanticRole.FOOTNOTE,
        BrandLayoutSemanticRole.HEADER,
        BrandLayoutSemanticRole.HEADER_IMAGE,
        BrandLayoutSemanticRole.NUMBER,
        BrandLayoutSemanticRole.REFERENCE,
    ),
)
def test_v4_auxiliary_roles_never_become_page_or_card_titles(
    semantic_role: BrandLayoutSemanticRole,
) -> None:
    pages = (
        BrandOcrLayoutPage(
            page_number=1,
            blocks=(
                BrandOcrLayoutBlock(
                    ordinal=0,
                    kind=BrandOcrBlockKind.TEXT,
                    text="辅助版面文字",
                    semantic_role=semantic_role,
                    normalized_bbox=(0.08, 0.10, 0.70, 0.16),
                ),
                BrandOcrLayoutBlock(
                    ordinal=1,
                    kind=BrandOcrBlockKind.TEXT,
                    text="相邻正文仍须独立。",
                    semantic_role=BrandLayoutSemanticRole.CONTENT,
                    normalized_bbox=(0.08, 0.18, 0.70, 0.32),
                ),
            ),
        ),
    )
    parser = _parser()
    parsed = parser.parse_ocr(
        markdown="compatibility only",
        layout_pages=pages,
        page_count=1,
    )
    chunking = parser.chunk(version_id=UUID(int=9), document=parsed)

    assert parsed.sections[0].title == "第 1 页"
    assert [chunk.text for chunk in chunking.chunks] == [
        "辅助版面文字",
        "相邻正文仍须独立。",
    ]


def test_v4_does_not_merge_ambiguous_columns() -> None:
    pages = (
        BrandOcrLayoutPage(
            page_number=1,
            blocks=(
                BrandOcrLayoutBlock(
                    ordinal=0,
                    kind=BrandOcrBlockKind.TEXT,
                    text="产品能力",
                    normalized_bbox=(0.05, 0.10, 0.35, 0.16),
                ),
                BrandOcrLayoutBlock(
                    ordinal=1,
                    kind=BrandOcrBlockKind.TEXT,
                    text="另一列中的独立说明。",
                    normalized_bbox=(0.60, 0.18, 0.95, 0.34),
                ),
            ),
        ),
    )
    parser = _parser()
    parsed = parser.parse_ocr(markdown="compatibility only", layout_pages=pages, page_count=1)
    result = parser.chunk(version_id=UUID(int=7), document=parsed)

    assert [chunk.text for chunk in result.chunks] == ["产品能力", "另一列中的独立说明。"]


def test_v4_fails_closed_when_layout_pages_are_missing() -> None:
    with pytest.raises(BrandUploadRejectedError) as raised:
        _parser().parse_ocr(markdown="# Markdown 不能替代布局", layout_pages=(), page_count=2)

    assert raised.value.code == "brand_ocr_layout_missing"


def test_v4_keeps_docx_source_structure_equivalent_to_v3() -> None:
    document = DocxDocument()
    document.add_heading("访谈", level=1)
    document.add_paragraph("Q1. 为什么重视探索\uff1f")
    document.add_paragraph("通过真实问题支持学习。")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "观察"
    table.cell(0, 1).text = "实践"
    buffer = BytesIO()
    document.save(buffer)
    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    v3 = _parser(layout=False).parse(body=buffer.getvalue(), media_type=media_type)
    v4 = _parser().parse(body=buffer.getvalue(), media_type=media_type)

    assert v4 == v3
