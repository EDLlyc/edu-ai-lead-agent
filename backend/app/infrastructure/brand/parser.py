from __future__ import annotations

import re
import zipfile
from dataclasses import dataclass
from io import BytesIO
from uuid import UUID, uuid5

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from app.core.errors import BrandUploadRejectedError
from app.domain.brand_knowledge import (
    LEGACY_BRAND_DERIVATION_VERSIONS,
    SUPPORTED_BRAND_DERIVATION_VERSIONS,
    BrandChunk,
    BrandChunkingResult,
    BrandClaimScope,
    BrandContentType,
    BrandDocumentKind,
    BrandSection,
    BrandSectionKind,
    ParsedBrandDocument,
    ParsedBrandSection,
    build_brand_embedding_text,
    classify_brand_chunk,
    normalize_brand_text,
)
from app.domain.value_objects import sha256_bytes, stable_key

_PARAGRAPH_BREAK_PATTERN = re.compile(r"\n[ \t]*\n(?:[ \t]*\n)*")
_MARKDOWN_BLOCK_BREAK_PATTERN = re.compile(
    r"\n(?=(?:#{1,6}[ \t]+|[-*+][ \t]+|\d+[.)][ \t]+|>[ \t]+))"
)
_SENTENCE_BREAK_PATTERN = re.compile(
    r"[\u3002\uff01\uff1f\uff1b!?;](?:[\"'\u201d\u2019\u300c\u300d\u300e\u300f\uff08\uff09()\uff3b\uff3d\u3010\u3011\u300a\u300b\u3008\u3009]*)"
)
_LINE_BREAK_PATTERN = re.compile(r"\n")
_PURE_PAGE_NUMBER_PATTERN = re.compile(r"^(?:[-\u2014\u2013]\s*)?\d{1,4}(?:\s*[-\u2014\u2013])?$")
_QUESTION_PATTERN = re.compile(
    r"^\s*(?:Q\s*(?P<arabic>\d+)|问题\s*(?P<chinese>[一二三四五六七八九十]+)|"
    r"第\s*(?P<ordinal>[一二三四五六七八九十\d]+)\s*问)"
    r"\s*[.\uff0e\u3001:\uff1a]?\s*(?P<question>.+?)\s*$",
    re.IGNORECASE,
)
_HEADING_PATTERN = re.compile(
    r"^(?:#{1,6}\s+|[一二三四五六七八九十]+、|第[一二三四五六七八九十\d]+(?:部分|章))"
)
_CHINESE_DIGITS = {
    "一": 1,
    "二": 2,
    "三": 3,
    "四": 4,
    "五": 5,
    "六": 6,
    "七": 7,
    "八": 8,
    "九": 9,
    "十": 10,
}
_DOCX_MAX_FILES = 2_000
_DOCX_MAX_EXPANDED_BYTES = 40 * 1024 * 1024
_DOCX_MAX_COMPRESSION_RATIO = 100


@dataclass(frozen=True, slots=True)
class _SectionDraft:
    kind: BrandSectionKind
    title: str
    text: str
    source_page: int | None = None
    question_number: int | None = None
    question_text: str | None = None


class BoundedBrandDocumentParser:
    def __init__(
        self,
        *,
        max_pages: int,
        max_characters: int,
        max_chunks: int,
        chunk_characters: int,
        overlap_characters: int,
        parser_version: str,
        chunk_version: str,
        embedding_input_version: str,
        sparse_text_threshold: int = 40,
    ) -> None:
        self._max_pages = max_pages
        self._max_characters = max_characters
        self._max_chunks = max_chunks
        self._chunk_characters = chunk_characters
        self._overlap_characters = overlap_characters
        self._parser_version = parser_version
        self._chunk_version = chunk_version
        self._embedding_input_version = embedding_input_version
        self._derivation_versions = (
            parser_version,
            chunk_version,
            embedding_input_version,
        )
        self._sparse_text_threshold = sparse_text_threshold
        if sparse_text_threshold < 1:
            raise ValueError("sparse text threshold must be positive")
        if self._derivation_versions not in SUPPORTED_BRAND_DERIVATION_VERSIONS:
            raise ValueError("unsupported brand parser/chunk/input version bundle")

    @property
    def _legacy_mode(self) -> bool:
        return self._derivation_versions == LEGACY_BRAND_DERIVATION_VERSIONS

    def parse(self, *, body: bytes, media_type: str) -> ParsedBrandDocument:
        if media_type == "application/pdf":
            parsed = self._parse_pdf(body)
        elif (
            media_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            parsed = self._parse_docx(body)
        elif media_type in {"text/plain", "text/markdown"}:
            decoded = self._decode_text(body)
            normalized = normalize_brand_text(decoded)
            if not normalized:
                raise BrandUploadRejectedError(
                    "brand_text_empty", "brand document contains no usable text"
                )
            parsed = self._assemble_document(
                drafts=(
                    _SectionDraft(
                        kind=BrandSectionKind.GENERIC,
                        title=self._text_title(normalized),
                        text=normalized,
                    ),
                ),
                page_count=None,
            )
        else:
            raise BrandUploadRejectedError(
                "unsupported_brand_file", "brand file type is unsupported"
            )
        if self._legacy_mode:
            normalized = normalize_brand_text(parsed.text)
            if len(normalized) > self._max_characters:
                raise BrandUploadRejectedError(
                    "brand_text_too_large", "parsed brand text exceeded the configured limit"
                )
            if not normalized and not parsed.requires_ocr:
                raise BrandUploadRejectedError(
                    "brand_text_empty", "brand document contains no usable text"
                )
            return ParsedBrandDocument(
                text=normalized,
                page_count=parsed.page_count,
                extraction_method=parsed.extraction_method,
                requires_ocr=parsed.requires_ocr,
                ocr_provider=parsed.ocr_provider,
                ocr_model=parsed.ocr_model,
                ocr_request_fingerprint=parsed.ocr_request_fingerprint,
                ocr_provider_request_id=parsed.ocr_provider_request_id,
                ocr_page_count=parsed.ocr_page_count,
                ocr_prompt_tokens=parsed.ocr_prompt_tokens,
                ocr_completion_tokens=parsed.ocr_completion_tokens,
                ocr_latency_ms=parsed.ocr_latency_ms,
            )
        # OCR adapters and focused tests may supply a document without section offsets.
        if parsed.text and not parsed.sections:
            normalized = normalize_brand_text(parsed.text)
            parsed = self._assemble_document(
                drafts=(
                    _SectionDraft(
                        kind=BrandSectionKind.GENERIC,
                        title=self._text_title(normalized),
                        text=normalized,
                    ),
                ),
                page_count=parsed.page_count,
                extraction_method=parsed.extraction_method,
                requires_ocr=parsed.requires_ocr,
            )
        if len(parsed.text) > self._max_characters:
            raise BrandUploadRejectedError(
                "brand_text_too_large", "parsed brand text exceeded the configured limit"
            )
        if not parsed.text and not parsed.requires_ocr:
            raise BrandUploadRejectedError(
                "brand_text_empty", "brand document contains no usable text"
            )
        return parsed

    def chunk(
        self,
        *,
        version_id: UUID,
        document: ParsedBrandDocument,
        document_title: str = "品牌资料",
        document_kind: BrandDocumentKind = BrandDocumentKind.OTHER,
    ) -> BrandChunkingResult:
        if self._legacy_mode:
            return self._chunk_legacy(version_id=version_id, document=document)
        parsed_sections = document.sections or self._legacy_sections(document)
        sections: list[BrandSection] = []
        chunks: list[BrandChunk] = []
        for parsed_section in parsed_sections:
            section_text_hash = sha256_bytes(parsed_section.text.encode("utf-8"))
            section_key = stable_key(
                version_id,
                parsed_section.ordinal,
                parsed_section.kind.value,
                section_text_hash,
                self._chunk_version,
            )
            section = BrandSection(
                id=uuid5(version_id, section_key),
                version_id=version_id,
                ordinal=parsed_section.ordinal,
                section_key=section_key,
                kind=parsed_section.kind,
                title=parsed_section.title,
                text=parsed_section.text,
                text_hash=section_text_hash,
                char_start=parsed_section.char_start,
                char_end=parsed_section.char_end,
                source_page=parsed_section.source_page,
                question_number=parsed_section.question_number,
                question_text=parsed_section.question_text,
            )
            sections.append(section)
            section_child_ordinal = 0
            child_spans = self._structured_child_spans(
                section=section,
                extraction_method=document.extraction_method,
                remaining_chunk_budget=self._max_chunks - len(chunks),
            )
            for chunk_local_start, chunk_local_end in child_spans:
                chunk_text = section.text[chunk_local_start:chunk_local_end]
                chunk_start = section.char_start + chunk_local_start
                chunk_end = section.char_start + chunk_local_end
                if document.text[chunk_start:chunk_end] != chunk_text:
                    raise ValueError("brand chunk must be an exact document slice")
                content_type, claim_scope, verification_required = classify_brand_chunk(
                    document_kind=document_kind,
                    section_title=section.title,
                    question_text=section.question_text,
                    text=chunk_text,
                )
                embedding_text = build_brand_embedding_text(
                    document_title=document_title,
                    section_title=section.title,
                    question_text=section.question_text,
                    content_type=content_type,
                    raw_text=chunk_text,
                )
                text_hash = sha256_bytes(chunk_text.encode("utf-8"))
                embedding_input_hash = sha256_bytes(embedding_text.encode("utf-8"))
                ordinal = len(chunks)
                chunk_key = stable_key(
                    version_id,
                    section.section_key,
                    section_child_ordinal,
                    text_hash,
                    embedding_input_hash,
                    self._chunk_version,
                )
                chunks.append(
                    BrandChunk(
                        id=uuid5(version_id, chunk_key),
                        section_id=section.id,
                        ordinal=ordinal,
                        section_ordinal=section_child_ordinal,
                        text=chunk_text,
                        text_hash=text_hash,
                        embedding_text=embedding_text,
                        embedding_input_hash=embedding_input_hash,
                        content_type=content_type,
                        claim_scope=claim_scope,
                        verification_required=verification_required,
                        char_start=chunk_start,
                        char_end=chunk_end,
                        chunk_key=chunk_key,
                    )
                )
                section_child_ordinal += 1
        if not chunks:
            raise BrandUploadRejectedError(
                "brand_text_empty", "brand document contains no usable text"
            )
        return BrandChunkingResult(sections=tuple(sections), chunks=tuple(chunks))

    def _chunk_legacy(
        self, *, version_id: UUID, document: ParsedBrandDocument
    ) -> BrandChunkingResult:
        chunks: list[BrandChunk] = []
        text = document.text
        start = 0
        while start < len(text):
            hard_end = min(start + self._chunk_characters, len(text))
            end = self._find_chunk_end(text=text, start=start, hard_end=hard_end)
            raw_chunk = text[start:end]
            leading_whitespace = len(raw_chunk) - len(raw_chunk.lstrip())
            trailing_whitespace = len(raw_chunk) - len(raw_chunk.rstrip())
            chunk_start = start + leading_whitespace
            chunk_end = end - trailing_whitespace
            chunk_text = text[chunk_start:chunk_end]
            if chunk_text:
                text_hash = sha256_bytes(chunk_text.encode("utf-8"))
                ordinal = len(chunks)
                chunk_key = stable_key(version_id, ordinal, text_hash, self._chunk_version)
                chunks.append(
                    BrandChunk(
                        id=uuid5(version_id, chunk_key),
                        section_id=None,
                        ordinal=ordinal,
                        section_ordinal=None,
                        text=chunk_text,
                        text_hash=text_hash,
                        embedding_text=chunk_text,
                        embedding_input_hash=text_hash,
                        content_type=BrandContentType.OTHER,
                        claim_scope=BrandClaimScope.BRAND_STATEMENT,
                        verification_required=False,
                        char_start=chunk_start,
                        char_end=chunk_end,
                        chunk_key=chunk_key,
                    )
                )
                if len(chunks) > self._max_chunks:
                    raise BrandUploadRejectedError(
                        "brand_chunk_limit", "brand document produced too many chunks"
                    )
            if end >= len(text):
                break
            next_start = max(start + 1, end - self._overlap_characters)
            while next_start < end and text[next_start].isspace():
                next_start += 1
            start = next_start
        if not chunks:
            raise BrandUploadRejectedError(
                "brand_text_empty", "brand document contains no usable text"
            )
        return BrandChunkingResult(sections=(), chunks=tuple(chunks))

    def _find_chunk_end(
        self, *, text: str, start: int, hard_end: int, text_end: int | None = None
    ) -> int:
        boundary = len(text) if text_end is None else text_end
        if hard_end >= boundary:
            return hard_end
        lower_bound = min(hard_end, start + max(1, int(self._chunk_characters * 0.6)))
        for pattern in (
            _PARAGRAPH_BREAK_PATTERN,
            _MARKDOWN_BLOCK_BREAK_PATTERN,
            _SENTENCE_BREAK_PATTERN,
            _LINE_BREAK_PATTERN,
        ):
            candidates = [match.end() for match in pattern.finditer(text, lower_bound, hard_end)]
            if candidates:
                return candidates[-1]
        return hard_end

    def _structured_child_spans(
        self,
        *,
        section: BrandSection,
        extraction_method: str,
        remaining_chunk_budget: int,
    ) -> tuple[tuple[int, int], ...]:
        preferred_ranges = self._preferred_child_ranges(section.text)
        spans = self._bounded_child_spans(section.text, preferred_ranges)
        if len(spans) <= remaining_chunk_budget:
            return spans
        if extraction_method != "ocr" or section.kind != BrandSectionKind.GENERIC:
            raise BrandUploadRejectedError(
                "brand_chunk_limit", "brand document produced too many chunks"
            )

        coalesced_ranges = self._coalesce_adjacent_ranges(preferred_ranges)
        spans = self._bounded_child_spans(section.text, coalesced_ranges)
        if len(spans) <= remaining_chunk_budget:
            return spans

        continuous_range = self._trimmed_parent_range(section.text)
        spans = self._bounded_child_spans(section.text, continuous_range)
        if len(spans) <= remaining_chunk_budget:
            return spans
        raise BrandUploadRejectedError(
            "brand_chunk_limit", "brand document produced too many chunks"
        )

    def _bounded_child_spans(
        self,
        text: str,
        ranges: tuple[tuple[int, int], ...],
    ) -> tuple[tuple[int, int], ...]:
        spans: list[tuple[int, int]] = []
        for range_start, range_end in ranges:
            local_start = range_start
            while local_start < range_end:
                hard_end = min(local_start + self._chunk_characters, range_end)
                local_end = self._find_chunk_end(
                    text=text,
                    start=local_start,
                    hard_end=hard_end,
                    text_end=range_end,
                )
                raw_chunk = text[local_start:local_end]
                leading_whitespace = len(raw_chunk) - len(raw_chunk.lstrip())
                trailing_whitespace = len(raw_chunk) - len(raw_chunk.rstrip())
                chunk_start = local_start + leading_whitespace
                chunk_end = local_end - trailing_whitespace
                if chunk_start < chunk_end:
                    spans.append((chunk_start, chunk_end))
                if local_end >= range_end:
                    break
                next_start = max(local_start + 1, local_end - self._overlap_characters)
                while next_start < local_end and text[next_start].isspace():
                    next_start += 1
                local_start = next_start
        return tuple(spans)

    def _coalesce_adjacent_ranges(
        self, ranges: tuple[tuple[int, int], ...]
    ) -> tuple[tuple[int, int], ...]:
        if len(ranges) < 2:
            return ranges
        coalesced: list[tuple[int, int]] = []
        current_start, current_end = ranges[0]
        for range_start, range_end in ranges[1:]:
            if range_end - current_start <= self._chunk_characters:
                current_end = range_end
                continue
            coalesced.append((current_start, current_end))
            current_start, current_end = range_start, range_end
        coalesced.append((current_start, current_end))
        return tuple(coalesced)

    @staticmethod
    def _trimmed_parent_range(text: str) -> tuple[tuple[int, int], ...]:
        start = len(text) - len(text.lstrip())
        end = len(text.rstrip())
        return ((start, end),) if start < end else ()

    @staticmethod
    def _preferred_child_ranges(text: str) -> tuple[tuple[int, int], ...]:
        """Return exact parent-local card/paragraph spans before bounded splitting."""

        boundaries = tuple(_PARAGRAPH_BREAK_PATTERN.finditer(text))
        if not boundaries:
            return ((0, len(text)),)
        ranges: list[tuple[int, int]] = []
        start = 0
        for boundary in boundaries:
            end = boundary.start()
            while start < end and text[start].isspace():
                start += 1
            while end > start and text[end - 1].isspace():
                end -= 1
            if start < end:
                ranges.append((start, end))
            start = boundary.end()
        end = len(text)
        while start < end and text[start].isspace():
            start += 1
        while end > start and text[end - 1].isspace():
            end -= 1
        if start < end:
            ranges.append((start, end))
        if len(ranges) < 2:
            return tuple(ranges) or ((0, len(text)),)
        first_start, first_end = ranges[0]
        first_text = text[first_start:first_end]
        first_is_question = _QUESTION_PATTERN.fullmatch(first_text.splitlines()[0]) is not None
        if len(first_text) <= 80 and (
            first_is_question or _SENTENCE_BREAK_PATTERN.search(first_text) is None
        ):
            _, second_end = ranges[1]
            ranges[:2] = [(first_start, second_end)]
        return tuple(ranges)

    def _parse_pdf(self, body: bytes) -> ParsedBrandDocument:
        if self._legacy_mode:
            return self._parse_pdf_legacy(body)
        try:
            reader = PdfReader(BytesIO(body), strict=True)
        except (PdfReadError, ValueError, TypeError, OSError):
            raise BrandUploadRejectedError(
                "malformed_brand_pdf", "brand PDF could not be parsed"
            ) from None
        if reader.is_encrypted:
            raise BrandUploadRejectedError(
                "encrypted_brand_pdf", "encrypted brand PDF is not accepted"
            )
        if len(reader.pages) > self._max_pages:
            raise BrandUploadRejectedError("brand_page_limit", "brand PDF exceeded the page limit")
        drafts: list[_SectionDraft] = []
        extracted_character_count = 0
        try:
            for page_number, page in enumerate(reader.pages, 1):
                raw_page_text = page.extract_text() or ""
                extracted_character_count += len(raw_page_text)
                if extracted_character_count > self._max_characters:
                    raise BrandUploadRejectedError(
                        "brand_text_too_large", "parsed brand text exceeded the configured limit"
                    )
                page_text = self._normalize_pdf_page(raw_page_text)
                if not page_text:
                    continue
                drafts.append(
                    _SectionDraft(
                        kind=BrandSectionKind.PAGE,
                        title=self._page_title(page_text, page_number),
                        text=page_text,
                        source_page=page_number,
                    )
                )
        except BrandUploadRejectedError:
            raise
        except Exception:
            raise BrandUploadRejectedError(
                "malformed_brand_pdf", "brand PDF could not be parsed"
            ) from None
        page_count = len(reader.pages)
        usable_character_count = sum(len(draft.text) for draft in drafts)
        requires_ocr = usable_character_count < page_count * self._sparse_text_threshold
        return self._assemble_document(
            drafts=tuple(drafts),
            page_count=page_count,
            requires_ocr=requires_ocr,
        )

    def _parse_pdf_legacy(self, body: bytes) -> ParsedBrandDocument:
        try:
            reader = PdfReader(BytesIO(body), strict=True)
        except (PdfReadError, ValueError, TypeError, OSError):
            raise BrandUploadRejectedError(
                "malformed_brand_pdf", "brand PDF could not be parsed"
            ) from None
        if reader.is_encrypted:
            raise BrandUploadRejectedError(
                "encrypted_brand_pdf", "encrypted brand PDF is not accepted"
            )
        if len(reader.pages) > self._max_pages:
            raise BrandUploadRejectedError("brand_page_limit", "brand PDF exceeded the page limit")
        parts: list[str] = []
        character_count = 0
        try:
            for page in reader.pages:
                page_text = page.extract_text() or ""
                character_count += len(page_text)
                if character_count > self._max_characters:
                    raise BrandUploadRejectedError(
                        "brand_text_too_large", "parsed brand text exceeded the configured limit"
                    )
                parts.append(page_text)
        except BrandUploadRejectedError:
            raise
        except Exception:
            raise BrandUploadRejectedError(
                "malformed_brand_pdf", "brand PDF could not be parsed"
            ) from None
        page_count = len(reader.pages)
        extracted_text = "\n\n".join(parts)
        requires_ocr = len(extracted_text.strip()) < page_count * self._sparse_text_threshold
        return ParsedBrandDocument(
            text=extracted_text,
            page_count=page_count,
            requires_ocr=requires_ocr,
        )

    def _parse_docx(self, body: bytes) -> ParsedBrandDocument:
        if self._legacy_mode:
            return self._parse_docx_legacy(body)
        self._inspect_docx_archive(body)
        try:
            document = Document(BytesIO(body))
        except (ValueError, TypeError, OSError, zipfile.BadZipFile):
            raise BrandUploadRejectedError(
                "malformed_brand_docx", "brand DOCX could not be parsed"
            ) from None
        blocks: list[tuple[str, bool]] = []
        try:
            for block in document.iter_inner_content():
                if isinstance(block, Paragraph):
                    text = normalize_brand_text(block.text)
                    if not text:
                        continue
                    style_name = (block.style.name if block.style is not None else "").casefold()
                    is_heading = style_name.startswith("heading") or style_name.startswith("标题")
                    blocks.append((text, is_heading or _HEADING_PATTERN.match(text) is not None))
                elif isinstance(block, Table):
                    rows: list[str] = []
                    for row in block.rows:
                        cells = [normalize_brand_text(cell.text) for cell in row.cells]
                        row_text = "\t".join(cell for cell in cells if cell)
                        if row_text:
                            rows.append(row_text)
                    if rows:
                        blocks.append(("\n".join(rows), False))
        except Exception:
            raise BrandUploadRejectedError(
                "malformed_brand_docx", "brand DOCX could not be parsed"
            ) from None
        if not blocks:
            raise BrandUploadRejectedError(
                "brand_text_empty", "brand document contains no usable text"
            )
        return self._assemble_document(
            drafts=self._docx_section_drafts(blocks),
            page_count=None,
        )

    def _parse_docx_legacy(self, body: bytes) -> ParsedBrandDocument:
        self._inspect_docx_archive(body)
        try:
            document = Document(BytesIO(body))
        except (ValueError, TypeError, OSError, zipfile.BadZipFile):
            raise BrandUploadRejectedError(
                "malformed_brand_docx", "brand DOCX could not be parsed"
            ) from None
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        if not parts:
            raise BrandUploadRejectedError(
                "brand_text_empty", "brand document contains no usable text"
            )
        return ParsedBrandDocument(text="\n\n".join(parts), page_count=None)

    def _docx_section_drafts(self, blocks: list[tuple[str, bool]]) -> tuple[_SectionDraft, ...]:
        drafts: list[_SectionDraft] = []
        current_kind = BrandSectionKind.GENERIC
        current_title = "文档内容"
        current_blocks: list[str] = []
        current_question_number: int | None = None
        current_question_text: str | None = None

        def flush() -> None:
            nonlocal current_blocks
            if not current_blocks:
                return
            drafts.append(
                _SectionDraft(
                    kind=current_kind,
                    title=current_title,
                    text="\n\n".join(current_blocks),
                    question_number=current_question_number,
                    question_text=current_question_text,
                )
            )
            current_blocks = []

        for block_text, is_heading in blocks:
            question = self._question_identity(block_text)
            if question is not None:
                flush()
                current_kind = BrandSectionKind.INTERVIEW_QA
                current_question_number, current_question_text = question
                current_title = current_question_text[:240]
                current_blocks = [block_text]
                continue
            if is_heading:
                flush()
                current_kind = BrandSectionKind.HEADING
                current_title = self._clean_heading(block_text)[:240]
                current_question_number = None
                current_question_text = None
                current_blocks = [block_text]
                continue
            current_blocks.append(block_text)
        flush()
        return tuple(drafts)

    def _assemble_document(
        self,
        *,
        drafts: tuple[_SectionDraft, ...],
        page_count: int | None,
        extraction_method: str = "local",
        requires_ocr: bool = False,
    ) -> ParsedBrandDocument:
        text_parts: list[str] = []
        sections: list[ParsedBrandSection] = []
        cursor = 0
        for draft in drafts:
            section_text = normalize_brand_text(draft.text)
            if not section_text:
                continue
            if text_parts:
                text_parts.append("\n\n")
                cursor += 2
            start = cursor
            text_parts.append(section_text)
            cursor += len(section_text)
            sections.append(
                ParsedBrandSection(
                    ordinal=len(sections),
                    kind=draft.kind,
                    title=normalize_brand_text(draft.title)[:240],
                    text=section_text,
                    char_start=start,
                    char_end=cursor,
                    source_page=draft.source_page,
                    question_number=draft.question_number,
                    question_text=draft.question_text,
                )
            )
        return ParsedBrandDocument(
            text="".join(text_parts),
            page_count=page_count,
            sections=tuple(sections),
            extraction_method=extraction_method,
            requires_ocr=requires_ocr,
        )

    @staticmethod
    def _legacy_sections(document: ParsedBrandDocument) -> tuple[ParsedBrandSection, ...]:
        leading = len(document.text) - len(document.text.lstrip())
        trailing = len(document.text) - len(document.text.rstrip())
        end = len(document.text) - trailing
        section_text = document.text[leading:end]
        if not section_text:
            return ()
        return (
            ParsedBrandSection(
                ordinal=0,
                kind=BrandSectionKind.GENERIC,
                title="文档内容",
                text=section_text,
                char_start=leading,
                char_end=end,
            ),
        )

    @staticmethod
    def _normalize_pdf_page(value: str) -> str:
        normalized = normalize_brand_text(value)
        lines = [
            line
            for line in normalized.splitlines()
            if _PURE_PAGE_NUMBER_PATTERN.fullmatch(line.strip()) is None
        ]
        return normalize_brand_text("\n".join(lines))

    @staticmethod
    def _page_title(text: str, page_number: int) -> str:
        first_line = next((line.strip() for line in text.splitlines() if line.strip()), "")
        if first_line and len(first_line) <= 80:
            return first_line
        return f"第 {page_number} 页"

    @staticmethod
    def _text_title(text: str) -> str:
        first_line = next((line.strip() for line in text.splitlines() if line.strip()), "")
        cleaned = BoundedBrandDocumentParser._clean_heading(first_line)
        return cleaned[:240] if cleaned and len(cleaned) <= 240 else "文档内容"

    @staticmethod
    def _clean_heading(value: str) -> str:
        return re.sub(r"^#{1,6}\s+", "", value).strip()

    @staticmethod
    def _question_identity(value: str) -> tuple[int, str] | None:
        first_line = value.splitlines()[0]
        match = _QUESTION_PATTERN.fullmatch(first_line)
        if match is None:
            return None
        number_text = match.group("arabic") or match.group("chinese") or match.group("ordinal")
        if number_text is None:
            return None
        number = int(number_text) if number_text.isdigit() else _parse_chinese_number(number_text)
        question = match.group("question").strip()
        if number < 1 or not question:
            return None
        return number, question

    @staticmethod
    def _inspect_docx_archive(body: bytes) -> None:
        try:
            with zipfile.ZipFile(BytesIO(body)) as archive:
                entries = archive.infolist()
                if len(entries) > _DOCX_MAX_FILES:
                    raise BrandUploadRejectedError(
                        "brand_archive_limit", "brand DOCX archive contains too many entries"
                    )
                expanded_bytes = 0
                for entry in entries:
                    name = entry.filename.casefold()
                    if "vbaproject.bin" in name or name.startswith("word/embeddings/"):
                        raise BrandUploadRejectedError(
                            "unsafe_brand_docx", "brand DOCX contains prohibited embedded content"
                        )
                    expanded_bytes += entry.file_size
                    if expanded_bytes > _DOCX_MAX_EXPANDED_BYTES:
                        raise BrandUploadRejectedError(
                            "brand_archive_limit", "brand DOCX expanded content is too large"
                        )
                    if (
                        entry.compress_size > 0
                        and entry.file_size / entry.compress_size > _DOCX_MAX_COMPRESSION_RATIO
                    ):
                        raise BrandUploadRejectedError(
                            "brand_archive_limit", "brand DOCX compression ratio is unsafe"
                        )
                    if name.endswith(".rels"):
                        relationships = archive.read(entry)
                        if (
                            b'TargetMode="External"' in relationships
                            or b"TargetMode='External'" in relationships
                        ):
                            raise BrandUploadRejectedError(
                                "unsafe_brand_docx", "brand DOCX contains external relationships"
                            )
        except BrandUploadRejectedError:
            raise
        except (zipfile.BadZipFile, KeyError, OSError):
            raise BrandUploadRejectedError(
                "malformed_brand_docx", "brand DOCX could not be parsed"
            ) from None

    @staticmethod
    def _decode_text(body: bytes) -> str:
        try:
            return body.decode("utf-8-sig")
        except UnicodeDecodeError:
            raise BrandUploadRejectedError(
                "invalid_brand_encoding", "brand text must be UTF-8"
            ) from None


def _parse_chinese_number(value: str) -> int:
    if value in _CHINESE_DIGITS:
        return _CHINESE_DIGITS[value]
    if value.startswith("十") and len(value) == 2:
        return 10 + _CHINESE_DIGITS.get(value[1], 0)
    if value.endswith("十") and len(value) == 2:
        return _CHINESE_DIGITS.get(value[0], 0) * 10
    if "十" in value and len(value) == 3:
        return _CHINESE_DIGITS.get(value[0], 0) * 10 + _CHINESE_DIGITS.get(value[2], 0)
    return 0
