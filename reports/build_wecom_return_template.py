from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "wecom-self-built-app-return-template.docx"

NAVY = "102A43"
OCEAN = "1769AA"
TEAL = "0F8B8D"
SKY = "EAF3FA"
MINT = "E6F4F1"
PAPER = "F7FAFC"
AMBER = "D97706"
RED = "B42318"
RED_PAPER = "FEF3F2"
LINE = "D9E2EC"
MUTED = "627D98"


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = LINE, size: str = "6") -> None:
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    properties = cell._tc.get_or_add_tcPr()
    width = properties.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        properties.append(width)
    width.set(qn("w:w"), str(int(width_cm * 567)))
    width.set(qn("w:type"), "dxa")


def set_run_font(run, *, size: float = 10.5, bold: bool = False, color: str = NAVY) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, *, before: float = 0, after: float = 4, line: float = 1.15) -> None:
    format_ = paragraph.paragraph_format
    format_.space_before = Pt(before)
    format_.space_after = Pt(after)
    format_.line_spacing = line


def add_text(paragraph, text: str, *, size: float = 10.5, bold: bool = False, color: str = NAVY):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return run


def add_labeled_paragraph(document: Document, label: str, text: str) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after=2)
    add_text(paragraph, f"{label}：", bold=True)
    add_text(paragraph, text, color="243B53")


def add_note(document: Document, title: str, text: str, *, fill: str = SKY, accent: str = OCEAN) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_width(cell, 17.0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=accent, size="10")
    set_cell_margins(cell, top=130, start=180, bottom=130, end=180)
    paragraph = cell.paragraphs[0]
    set_paragraph_spacing(paragraph, after=0)
    add_text(paragraph, f"{title}：", bold=True, color=accent)
    add_text(paragraph, text, color="243B53")
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_section_heading(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    set_paragraph_spacing(paragraph, before=10, after=4)
    add_text(paragraph, title, size=14, bold=True, color=OCEAN)
    ppr = paragraph._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), OCEAN)
    border.append(bottom)
    ppr.append(border)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_width(cell, widths[index])
        set_cell_shading(cell, NAVY)
        set_cell_border(cell, color="FFFFFF", size="6")
        set_cell_margins(cell, top=110, start=120, bottom=110, end=120)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        set_paragraph_spacing(paragraph, after=0)
        add_text(paragraph, header, size=9.5, bold=True, color="FFFFFF")

    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cell = cells[index]
            set_cell_width(cell, widths[index])
            set_cell_shading(cell, PAPER if row_index % 2 else "FFFFFF")
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            set_paragraph_spacing(paragraph, after=0)
            if value == "私密渠道":
                add_text(paragraph, value, size=9.5, bold=True, color=RED)
            else:
                add_text(paragraph, value, size=9.5, color="243B53")
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_checkbox_line(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after=3)
    add_text(paragraph, "□  ", size=11, color=OCEAN)
    add_text(paragraph, text, size=10.5, color="243B53")


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("243B53")

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(header, after=2)
    add_text(header, "企业微信参数回传模板", size=9, color=MUTED)
    header.add_run(" " * 4)
    add_text(header, "赛先生品牌内容系统", size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(footer, after=0)
    add_text(footer, "版本 v1.0 · 仅用于企业微信文字/图片投递 · 不包含朋友圈自动发布", size=8.5, color=MUTED)


def build_document() -> None:
    document = Document()
    configure_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, after=2)
    add_text(title, "赛先生科学内容系统", size=12, bold=True, color=OCEAN)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, after=3)
    add_text(title, "企业微信自建应用参数回传模板", size=21, bold=True, color=NAVY)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle, after=10)
    add_text(subtitle, "请范老师填写非敏感信息；Secret 请通过私密渠道单独提供", size=10.5, color=MUTED)

    add_note(
        document,
        "填写说明",
        "创建自建应用并设置可见范围后，填写下表。目标销售必须在应用可见范围内；userid 不是手机号。"
        "本系统第一阶段只向企业微信内部员工发送文字和图片。",
        fill=MINT,
        accent=TEAL,
    )

    add_section_heading(document, "一、企业与应用信息")
    add_table(
        document,
        ["字段", "填写内容", "说明", "敏感级别"],
        [
            ["企业名称", "（请填写）", "用于确认企业", "普通"],
            ["应用名称", "（建议：赛先生科学内容助手）", "企业微信后台自建应用名称", "普通"],
            ["WECOM_CORP_ID", "（请填写）", "“我的企业”/企业信息中的企业 ID", "内部信息"],
            ["WECOM_AGENT_ID", "（请填写）", "自建应用详情中的 AgentId，整数", "内部信息"],
            ["WECOM_CORP_SECRET", "请走私密渠道", "应用详情中的 Secret，不要写在普通聊天或群聊", "私密渠道"],
        ],
        [4.0, 4.1, 5.2, 3.7],
    )

    add_section_heading(document, "二、目标接收人信息")
    add_table(
        document,
        ["字段", "填写内容", "说明"],
        [
            ["目标销售姓名", "（请填写）", "例如：张三（销售）"],
            ["WECOM_DEFAULT_RECIPIENT_ID", "（请填写）", "通讯录成员的账号/userid，不要填手机号"],
            ["应用可见范围", "□ 已包含目标销售", "请在企业微信应用详情中确认"],
            ["首次测试接收人", "（姓名 + userid）", "首次只向这一名员工发送测试消息"],
        ],
        [5.0, 5.4, 6.6],
    )

    add_section_heading(document, "三、投递设置")
    add_table(
        document,
        ["设置项", "选择或填写", "建议"],
        [
            ["是否先做测试投递", "□ 是  □ 否", "建议先选择“是”"],
            ["是否开启每日正式自动投递", "□ 是  □ 否", "首次建议“否”，测试确认后再开启"],
            ["WECOM_DEFAULT_RECIPIENT_NAME", "（请填写）", "系统显示用的目标销售姓名"],
            ["后台是否要求可信 IP", "□ 是  □ 否", "若是，请补充服务器公网出口 IP"],
            ["服务器公网出口 IP", "（如需要，请填写）", "仅在企业微信后台提示时提供"],
        ],
        [6.0, 5.0, 6.0],
    )

    add_note(
        document,
        "Secret 安全要求",
        "WECOM_CORP_SECRET 不要发送到微信群、普通聊天、邮件正文、截图或 GitHub。"
        "请使用密码管理器、一次性安全链接、电话口述，或直接在服务器权限为 600 的 .env 中录入。",
        fill=RED_PAPER,
        accent=RED,
    )

    add_section_heading(document, "回传前确认")
    add_checkbox_line(document, "目标销售已经加入自建应用可见范围")
    add_checkbox_line(document, "已确认提供的是企业微信 userid/账号，而不是手机号")
    add_checkbox_line(document, "WECOM_CORP_ID、AgentId 和目标员工属于同一企业")
    add_checkbox_line(document, "Secret 将通过私密渠道单独提供")

    add_note(
        document,
        "请回传给项目负责人",
        "完成本模板后，可通过私聊发送企业名称、应用名称、CorpID、AgentId、目标销售姓名和 userid。"
        "Secret 请另外使用私密渠道提供；我们收到后再进行企业微信测试投递。",
        fill=SKY,
        accent=OCEAN,
    )

    document.save(OUTPUT)


if __name__ == "__main__":
    build_document()
